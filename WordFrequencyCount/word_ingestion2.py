import docx
import pprint as pp


def main():
    
    doc = docx.Document('word_documents\\2020 05 04 NR Daily Log.docx')

#    docaslist = list()
#
#    for table in doc.tables:
#        for row in table.rows:
#            for cell in row.cells:
#                docaslist.append(cell.text)
#
#    pp.pprint(docaslist)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '02' in paragraph.text:
                        print('found date: ', paragraph.text)

    
if __name__ == '__main__':
    main()
